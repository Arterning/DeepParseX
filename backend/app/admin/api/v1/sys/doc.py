#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from typing import Annotated, List

from fastapi import APIRouter, Depends, File, Form, Path, Query, Request, Body, UploadFile
from fastapi.responses import StreamingResponse
from fastapi.exceptions import HTTPException

from backend.core.conf import settings
from backend.app.admin.schema.doc import CreateSysDocParam, CollectDocParam, GetSysDocListDetails, GetSysDocPage, UpdateSysDocParam, GetDocDetail, ParseEntityParams, ExtractEntitiesParams, TranslateChunksParams, UpdateOcrPageParam, UpdateOcrPageTranslationParam, AiSearchParam, AiSearchResult, BatchMoveDocParam, AiOverviewParam, AiOverviewResult
from backend.app.admin.schema.doc_chunk import UpdateChunkContentParam
from backend.app.admin.service.doc_service import sys_doc_service
from backend.app.admin.service.upload_service import upload_service
from backend.common.pagination import DependsPagination, paging_data
from backend.common.response.response_schema import ResponseModel, response_base
from backend.common.security.jwt import DependsJwtAuth
from backend.common.security.permission import RequestPermission
from backend.common.security.rbac import DependsRBAC
from backend.database.db_pg import CurrentSession, async_db_session
from backend.utils.serializers import select_as_dict
from backend.utils.oss_client import minio_client

from minio.error import S3Error

router = APIRouter()


# collect_doc 收藏文件
@router.post('/collect', summary='收藏/取消收藏文件',
    dependencies=[DependsJwtAuth]
)
async def collect_doc(request: Request, obj: CollectDocParam) -> ResponseModel:
    user_id = request.user.id
    doc = await sys_doc_service.get(pk=obj.doc_id)
    if not doc:
        return response_base.fail(message='文件不存在')
    
    await sys_doc_service.collect_doc(user_id=user_id, collecton_id=obj.collection_id, doc_id=obj.doc_id)
    return response_base.success(data='操作成功')


# 构建文件的知识图谱
@router.post('/build_graph/{pk}', summary='构建文件的知识图谱',
    tags=['文档管理'],
    dependencies=[DependsJwtAuth]
 )
async def build_graph(
    pk: Annotated[int, Path(...)],
    obj: ParseEntityParams
) -> ResponseModel:
    await sys_doc_service.build_graph(pk=pk, entity_types=obj.entity_types)
    doc = await sys_doc_service.get(pk=pk)
    if not doc.doc_spos:
        return response_base.fail()
    triples = doc.doc_spos
    visualize_knowledge_graph = sys_doc_service.build_visualize_knowledge_graph(triples=triples)
    return response_base.success(data=visualize_knowledge_graph)


# 根据实体类型提取实体
@router.post('/extract_entities/{pk}', summary='根据实体类型提取实体',
    tags=['文档管理'],
    dependencies=[DependsJwtAuth]
)
async def extract_entities(
    pk: Annotated[int, Path(...)],
    obj: ExtractEntitiesParams
) -> ResponseModel:
    try:
        from backend.app.admin.crud.crud_entity_type import entity_type_dao

        type_definitions = []
        async with async_db_session() as db:
            for type_id in obj.entity_type_ids:
                entity_type = await entity_type_dao.get(db, type_id)
                if entity_type:
                    type_definitions.append({
                        "type_name": entity_type.name,
                        "description": entity_type.description or "",
                        "fields": entity_type.field_definition or []
                    })

        if not type_definitions:
            return response_base.fail(data="未找到指定的实体类型")

        entity_count = await sys_doc_service.extract_entities_by_types(
            pk=pk,
            type_definitions=type_definitions
        )
        return response_base.success(data={"count": entity_count } )
    except Exception as e:
        return response_base.fail(data=str(e))


# 提取内容
@router.get('/extract_text/{pk}', summary='提取文本',
    dependencies=[DependsJwtAuth]
 )
async def extract_text(pk: Annotated[int, Path(...)]) -> ResponseModel:
    data = await upload_service.extract_text(pk=pk)
    return response_base.success(data=data)


@router.get('/recent_docs', summary='获取最新上传文件',
    dependencies=[DependsJwtAuth]
 )
async def get_recent_docs(request: Request) -> ResponseModel:
    user_id = None if request.user.is_superuser else request.user.id
    docs = await sys_doc_service.get_hot_docs(user_id)
    hot_docs = [GetSysDocPage(**select_as_dict(doc)) for doc in docs]
    return response_base.success(data=hot_docs)


bucket_name = settings.BUCKET_NAME

# 获取原文件
@router.get("/preview/{obj_name}", summary = "预览文件")
async def preview_pdf(obj_name: str):
    try:
        # 从 MinIO 获取对象
        response = minio_client.get_object(bucket_name, obj_name)

        # 获取文件的 MIME 类型
        media_type = response.getheader('Content-Type')

        if media_type == 'application/msword':
            import tempfile
            import asyncio
            import os
            import subprocess

            try:
                # 创建临时文件
                with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as temp_doc:
                    
                    temp_doc.write(response.read())
                    temp_doc_path = temp_doc.name

            
                # 使用LibreOffice进行转换
                loop = asyncio.get_running_loop()
                await loop.run_in_executor(
                    None, 
                    lambda: subprocess.run([
                        "libreoffice",
                        "--headless",
                        "--convert-to", "pdf",
                        "--outdir", os.path.dirname(temp_doc_path),
                        temp_doc_path
                    ], check=True)
                )

                temp_pdf_path = os.path.splitext(temp_doc_path)[0] + ".pdf"

                # 读取PDF文件并创建生成器
                def pdf_generator():
                    try:
                        
                        with open(temp_pdf_path, "rb") as f:
                            while True:
                                chunk = f.read(9024)
                                if not chunk:
                                    break
                                yield chunk
                    finally:
                        # 清理临时文件
                        if os.path.exists(temp_doc_path):
                            os.unlink(temp_doc_path)
                        if os.path.exists(temp_pdf_path):
                            os.unlink(temp_pdf_path)

                # 返回PDF流
                return StreamingResponse(
                    pdf_generator(),
                    media_type="application/pdf",
                )
            except Exception as e:
                # 确保资源被释放
                if 'response' in locals():
                    response.close()
                    response.release_conn()
                if 'temp_doc_path' in locals() and os.path.exists(temp_doc_path):
                    os.unlink(temp_doc_path)
                if 'temp_pdf_path' in locals() and os.path.exists(temp_pdf_path):
                    os.unlink(temp_pdf_path)
                raise e
        
        
        async def file_generator(response):
            while True:
                chunk = response.read(9024)  # 逐块读取文件
                if not chunk:
                    break
                yield chunk
            response.close()
            response.release_conn()
        
        return StreamingResponse(file_generator(response), media_type=media_type)
    except S3Error as e:
        raise HTTPException(status_code=404, detail="File not found")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"服务器错误: {str(e)}")



@router.post('/convert_to_pdf/{pk}', summary='将文件转换为PDF（支持图片、Word、PPT）',
    dependencies=[DependsJwtAuth]
)
async def convert_to_pdf(
    pk: Annotated[int, Path(...)],
    request: Request,
) -> ResponseModel:
    import uuid as uuid_lib
    import asyncio
    from io import BytesIO

    doc = await sys_doc_service.get(pk=pk)
    if not doc:
        return response_base.fail(message='文件不存在')
    if doc.type not in ('image', 'document', 'ppt'):
        return response_base.fail(message='仅支持图片、Word文档、PPT文件转换')

    try:
        # 从 MinIO 获取原始文件
        minio_resp = minio_client.get_object(bucket_name, doc.file)
        file_bytes = minio_resp.read()
        minio_resp.close()
        minio_resp.release_conn()

        if doc.type == 'image':
            # 图片转PDF：Pillow
            from PIL import Image
            img = Image.open(BytesIO(file_bytes))
            if img.mode in ('RGBA', 'LA', 'P'):
                img = img.convert('RGB')
            pdf_buf = BytesIO()
            img.save(pdf_buf, format='PDF')
            pdf_bytes = pdf_buf.getvalue()

        else:
            # Word / PPT 转PDF：LibreOffice
            import tempfile
            import subprocess
            import os

            file_suffix = doc.file_suffix or ('.docx' if doc.type == 'document' else '.pptx')
            tmp_path = None
            tmp_pdf_path = None
            try:
                with tempfile.NamedTemporaryFile(suffix=file_suffix, delete=False) as tmp:
                    tmp.write(file_bytes)
                    tmp_path = tmp.name
                tmp_pdf_path = os.path.splitext(tmp_path)[0] + '.pdf'

                await asyncio.to_thread(
                    lambda: subprocess.run(
                        [
                            'libreoffice', '--headless',
                            '--convert-to', 'pdf',
                            '--outdir', os.path.dirname(tmp_path),
                            tmp_path,
                        ],
                        check=True,
                        capture_output=True,
                    )
                )

                with open(tmp_pdf_path, 'rb') as f:
                    pdf_bytes = f.read()
            finally:
                if tmp_path and os.path.exists(tmp_path):
                    os.unlink(tmp_path)
                if tmp_pdf_path and os.path.exists(tmp_pdf_path):
                    os.unlink(tmp_pdf_path)

        # 上传 PDF 到 MinIO
        unique_id = str(uuid_lib.uuid4())
        new_filename = f"{unique_id}.pdf"

        def _sync_put():
            minio_client.put_object(
                bucket_name, new_filename,
                BytesIO(pdf_bytes), len(pdf_bytes),
                'application/pdf',
            )
        await asyncio.to_thread(_sync_put)

        original_name = doc.title or doc.name or 'file'
        base_name = original_name.rsplit('.', 1)[0] if '.' in original_name else original_name
        pdf_title = f"{base_name}.pdf"

        from backend.app.admin.schema.doc import CreateSysDocParam
        obj = CreateSysDocParam(
            title=pdf_title,
            name=pdf_title,
            type='pdf',
            file=new_filename,
            uuid=unique_id,
            file_suffix='.pdf',
            size=len(pdf_bytes),
            status=0,
            doc_dir_id=doc.doc_dir_id,
            created_by=request.user.id,
            created_user=request.user.username,
        )
        new_doc = await sys_doc_service.create(obj=obj)
        return response_base.success(data={'id': new_doc.id, 'title': pdf_title})
    except S3Error:
        raise HTTPException(status_code=404, detail='原文件未找到')
    except Exception as e:
        raise HTTPException(status_code=500, detail=f'转换失败: {str(e)}')


@router.post('/convert_file_to_pdf', summary='将Word/PPT文件转换为PDF（不入库）', dependencies=[DependsJwtAuth])
async def convert_file_to_pdf(
    file: Annotated[UploadFile, File(description='Word 或 PPT 文件')],
) -> StreamingResponse:
    import tempfile
    import subprocess
    import os
    import asyncio
    from io import BytesIO
    from urllib.parse import quote

    filename = file.filename or 'document'
    dot = filename.rfind('.')
    suffix = filename[dot:].lower() if dot != -1 else ''

    SUPPORTED = {'.docx', '.doc', '.pptx', '.ppt'}
    if suffix not in SUPPORTED:
        raise HTTPException(status_code=400, detail=f'不支持的格式 {suffix}，仅支持 .docx / .doc / .pptx / .ppt')

    file_bytes = await file.read()

    tmp_path = None
    tmp_pdf_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        tmp_pdf_path = os.path.splitext(tmp_path)[0] + '.pdf'

        await asyncio.to_thread(
            lambda: subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'pdf',
                 '--outdir', os.path.dirname(tmp_path), tmp_path],
                check=True,
                capture_output=True,
            )
        )

        with open(tmp_pdf_path, 'rb') as f:
            pdf_bytes = f.read()
    except subprocess.CalledProcessError as e:
        raise HTTPException(status_code=500, detail=f'转换失败: {e.stderr.decode(errors="ignore")}')
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        if tmp_pdf_path and os.path.exists(tmp_pdf_path):
            os.unlink(tmp_pdf_path)

    stem = filename[:dot] if dot != -1 else filename
    pdf_name = stem + '.pdf'
    encoded_name = quote(pdf_name, safe='')

    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type='application/pdf',
        headers={'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_name}"},
    )


def _md_to_docx_bytes(content: str) -> bytes:
    from io import BytesIO
    import re
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for line in content.splitlines():
        # 标题
        h_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()
            doc.add_heading(text, level=level)
            continue

        # 分隔线
        if re.match(r'^[-*_]{3,}\s*$', line):
            doc.add_paragraph('─' * 40)
            continue

        # 无序列表
        ul_match = re.match(r'^[\*\-\+]\s+(.*)', line)
        if ul_match:
            p = doc.add_paragraph(style='List Bullet')
            _apply_inline(p, ul_match.group(1))
            continue

        # 有序列表
        ol_match = re.match(r'^\d+\.\s+(.*)', line)
        if ol_match:
            p = doc.add_paragraph(style='List Number')
            _apply_inline(p, ol_match.group(1))
            continue

        # 代码块标记行（``` 不渲染为段落）
        if re.match(r'^```', line):
            continue

        # 普通段落
        p = doc.add_paragraph()
        _apply_inline(p, line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.post('/convert_md_to_docx', summary='将 Markdown 文件转换为 Word（.docx）', dependencies=[DependsJwtAuth])
async def convert_md_to_docx(
    file: Annotated[UploadFile, File(description='Markdown 文件（.md）')],
) -> StreamingResponse:
    from io import BytesIO
    from urllib.parse import quote

    filename = file.filename or 'document.md'
    if not filename.lower().endswith('.md'):
        raise HTTPException(status_code=400, detail='仅支持 .md 格式')

    content = (await file.read()).decode('utf-8', errors='replace')
    docx_bytes = _md_to_docx_bytes(content)

    stem = filename[:-3] if filename.lower().endswith('.md') else filename
    safe_name = stem + '.docx'
    encoded_name = quote(safe_name, safe='')

    return StreamingResponse(
        BytesIO(docx_bytes),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_name}"},
    )


@router.post('/convert_md_text_to_docx', summary='将 Markdown 文本转换为 Word（.docx）', dependencies=[DependsJwtAuth])
async def convert_md_text_to_docx(
    content: Annotated[str, Body(..., description='Markdown 内容')],
    filename: Annotated[str, Body(description='输出文件名（不含扩展名）')] = 'document',
) -> StreamingResponse:
    from io import BytesIO
    from urllib.parse import quote

    docx_bytes = _md_to_docx_bytes(content)
    safe_name = (filename or 'document') + '.docx'
    encoded_name = quote(safe_name, safe='')

    return StreamingResponse(
        BytesIO(docx_bytes),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_name}"},
    )


def _apply_inline(paragraph, text: str):
    """解析行内 **bold**、*italic*、`code` 并添加 run。"""
    import re
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|(.+?)(?=\*\*|\*|`|$))', re.DOTALL)
    for m in pattern.finditer(text):
        if m.group(2):
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3):
            run = paragraph.add_run(m.group(3))
            run.italic = True
        elif m.group(4):
            run = paragraph.add_run(m.group(4))
            run.font.name = 'Courier New'
        elif m.group(5):
            paragraph.add_run(m.group(5))


@router.post('/merge_images_to_pdf', summary='合并多张图片为一个PDF文件', dependencies=[DependsJwtAuth])
async def merge_images_to_pdf(
    files: Annotated[List[UploadFile], File(description='图片文件列表，按顺序合并为PDF页面')],
    filename: Annotated[str, Form(description='合并后PDF的文件名（不含扩展名）')] = 'merged',
) -> StreamingResponse:
    from PIL import Image
    from io import BytesIO
    from urllib.parse import quote

    if not files:
        raise HTTPException(status_code=400, detail='未上传任何图片')

    pil_images = []
    for f in files:
        data = await f.read()
        try:
            img = Image.open(BytesIO(data))
            if img.mode not in ('RGB', 'L'):
                img = img.convert('RGB')
            pil_images.append(img)
        except Exception:
            raise HTTPException(status_code=400, detail=f'无法解析图片: {f.filename}')

    buf = BytesIO()
    pil_images[0].save(buf, format='PDF', save_all=True, append_images=pil_images[1:], quality=95)
    pdf_bytes = buf.getvalue()

    safe_name = (filename.strip() or 'merged') + '.pdf'
    encoded_name = quote(safe_name, safe='')

    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type='application/pdf',
        headers={'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_name}"},
    )


@router.post(
    '/ai_search',
    summary='AI 混合检索文件（向量+全文+RRF，无 topK 限制）',
    dependencies=[DependsJwtAuth],
)
async def ai_search_docs(obj: AiSearchParam) -> ResponseModel:
    from backend.app.admin.service.query_rewrite_service import rewrite_query
    rewritten = await rewrite_query(obj.question)
    items = await sys_doc_service.ai_search_docs(
        question=obj.question,
        rewritten_query=rewritten,
    )
    result = AiSearchResult(items=items, rewritten_query=rewritten, total=len(items))
    return response_base.success(data=result.model_dump())


@router.post(
    '/ai_overview',
    summary='AI 概览（基于已有检索结果生成 RAG 回答）',
    dependencies=[DependsJwtAuth],
)
async def ai_overview(obj: AiOverviewParam) -> ResponseModel:
    result = await sys_doc_service.ai_overview(
        question=obj.question,
        search_results=[item.model_dump() for item in obj.search_results],
    )
    return response_base.success(data=AiOverviewResult(**result).model_dump())


@router.put(
    '/batch_move',
    summary='批量移动文件到指定目录',
    dependencies=[DependsJwtAuth],
)
async def batch_move_docs(obj: BatchMoveDocParam) -> ResponseModel:
    count = await sys_doc_service.batch_move(doc_ids=obj.doc_ids, doc_dir_id=obj.doc_dir_id)
    return response_base.success(data={'moved': count})


@router.get(
    '/search',
    summary='（模糊条件）获取所有文件',
    dependencies=[
        DependsJwtAuth,
        DependsPagination,
    ],
)
async def search(
    tokens: Annotated[str | None, Query()] = None,
    keyword: Annotated[str | None, Query()] = None,
    page: Annotated[int | None, Query()] = None,
    size: Annotated[int | None, Query()] = None
) -> ResponseModel:
    keyword = keyword or tokens
    docs = await sys_doc_service.search(keyword=keyword, page=page, size=size)
    return response_base.success(data=docs)


@router.get(
    '/similar_search',
    summary='（模糊条件）获取所有文件',
    dependencies=[
        DependsJwtAuth,
        DependsPagination,
    ],
)
async def similar_search(
    query: Annotated[str | None, Query()] = None,
    page: Annotated[int | None, Query()] = None,
    size: Annotated[int | None, Query()] = None
) -> ResponseModel:
    data = await sys_doc_service.similar_search(query=query, page=page, size=size)
    return response_base.success(data=data)



@router.post('/extract_summary/{pk}', summary='提取摘要',
    dependencies=[DependsJwtAuth]
)
async def extract_summary(
    pk: Annotated[int, Path(...)],
) -> ResponseModel:
    data = await sys_doc_service.generate_summary(id=pk)
    return response_base.success(data=data)



@router.post('/translate_chunks/{pk}', summary='翻译文件所有分块',
    dependencies=[DependsJwtAuth]
)
async def translate_chunks(
    pk: Annotated[int, Path(...)],
    obj: TranslateChunksParams
) -> ResponseModel:
    doc = await sys_doc_service.get(pk=pk)
    if not doc:
        return response_base.fail(message='文件不存在')
    results = await sys_doc_service.translate_chunks(pk=pk, target_language=obj.target_language)
    return response_base.success(data=results)


@router.post('/translate_pages/{pk}', summary='翻译OCR分页',
    dependencies=[DependsJwtAuth]
)
async def translate_pages(
    pk: Annotated[int, Path(...)],
    obj: TranslateChunksParams
) -> ResponseModel:
    doc = await sys_doc_service.get(pk=pk)
    if not doc:
        return response_base.fail(message='文件不存在')
    results = await sys_doc_service.translate_pages(pk=pk, target_language=obj.target_language)
    return response_base.success(data=results)


@router.put('/chunk/{chunk_id}', summary='更新分块内容',
    dependencies=[DependsJwtAuth]
)
async def update_chunk(
    request: Request,
    chunk_id: Annotated[int, Path(...)],
    obj: UpdateChunkContentParam
) -> ResponseModel:
    result = await sys_doc_service.update_chunk(
        chunk_id=chunk_id,
        chunk_text=obj.chunk_text,
        chunk_translation=obj.chunk_translation,
        user_id=request.user.id,
        username=request.user.username,
    )
    return response_base.success(data=result)


@router.post('/compose_refined_markdown/{pk}', summary='生成精炼结构化Markdown（基于OCR翻译结果）',
    dependencies=[DependsJwtAuth]
)
async def compose_refined_markdown(
    pk: Annotated[int, Path(...)],
    obj: TranslateChunksParams
) -> ResponseModel:
    doc = await sys_doc_service.get(pk=pk)
    if not doc:
        return response_base.fail(message='文件不存在')

    # 如果尚未翻译，先翻译
    if not doc.ocr_pages_translation:
        await sys_doc_service.translate_pages(pk=pk, target_language=obj.target_language)

    result = await sys_doc_service.compose_refined_markdown(pk=pk)
    return response_base.success(data=result)


@router.put('/page/{pk}', summary='更新OCR分页内容',
    dependencies=[DependsJwtAuth]
)
async def update_ocr_page(
    pk: Annotated[int, Path(...)],
    request: Request,
    obj: UpdateOcrPageParam
) -> ResponseModel:
    result = await sys_doc_service.update_ocr_page(
        pk=pk,
        page=obj.page,
        text=obj.text,
        user_id=request.user.id,
        username=request.user.username,
    )
    return response_base.success(data=result)


@router.put('/page/{pk}/translation', summary='更新OCR分页翻译内容',
    dependencies=[DependsJwtAuth]
)
async def update_ocr_page_translation(
    pk: Annotated[int, Path(...)],
    request: Request,
    obj: UpdateOcrPageTranslationParam
) -> ResponseModel:
    result = await sys_doc_service.update_ocr_page_translation(
        pk=pk,
        page=obj.page,
        translation=obj.translation,
        user_id=request.user.id,
        username=request.user.username,
    )
    return response_base.success(data=result)


@router.get('/{pk}', summary='获取文件详情', dependencies=[DependsJwtAuth])
async def get_sys_doc(pk: Annotated[int, Path(...)], request: Request) -> ResponseModel:
    doc = await sys_doc_service.get(pk=pk)
    doc_data = []
    for data in doc.doc_data:
        doc_data.append(data.row)
    doc_dict = select_as_dict(doc)
    graph_data = sys_doc_service.build_visualize_knowledge_graph(triples=doc.doc_spos)

    doc_chunks = sorted(doc.doc_chunks, key=lambda c: c.chunk_index)
    doc_dict.update({"doc_data": doc_data})
    doc_dict.update({"doc_chunks": doc_chunks})
    doc_dict.update({"graph_data": graph_data})

    # 查询子文件（belong = 当前文件id的文件）
    children = await sys_doc_service.get_children(pk=pk)
    doc_dict.update({"children": children})

    # 查询该文档所在的所有收藏夹 ID
    user_id = request.user.id
    starred_ids = await sys_doc_service.get_doc_starred_ids(doc_id=pk, user_id=user_id)
    doc_dict.update({"starred_ids": starred_ids})

    data = GetDocDetail(**doc_dict)
    return response_base.success(data=data)


@router.get(
    '',
    summary='（模糊条件）分页获取所有文件',
    dependencies=[
        DependsJwtAuth,
        DependsPagination,
    ],
)
async def get_pagination_sys_doc(request: Request,
                                 db: CurrentSession,
                                 name: Annotated[str | None, Query()] = None,
                                 title: Annotated[str | None, Query()] = None,
                                 doc_type: Annotated[list[str] | None, Query()] = None,
                                 content: Annotated[str | None, Query()] = None,
                                 source: Annotated[str | None, Query()] = None,
                                 rangeValue: Annotated[list[str] | None, Query()] = ['', ''],
                                 tag_ids: Annotated[list[int] | None, Query()] = None,
                                 doc_dir_id: Annotated[int | None, Query()] = None,
                                 status: Annotated[int | None, Query()] = None,
                                ) -> ResponseModel:

    # 从JWT中获取当前登录用户ID
    # 超级管理员可查看所有文件，普通用户只能查看自己的文件
    current_user_id = None if request.user.is_superuser else request.user.id

    sys_doc_select = await sys_doc_service.get_select(
        name=name,
        title=title,
        doc_type=doc_type,
        source=source,
        content=content,
        rangeValue=rangeValue,
        current_user_id=current_user_id,
        tag_ids=tag_ids,
        doc_dir_id=doc_dir_id,
        status=status,
    )
    page_data = await paging_data(db, sys_doc_select, GetSysDocPage)

    # Add is_collected flag
    doc_ids = [item.get("id") for item in page_data['items']]
    if doc_ids:
        user_id = request.user.id
        collected_doc_ids = await sys_doc_service.get_collected_doc_ids(user_id, doc_ids)
        for item in page_data['items']:
            item["is_collected"] = item.get("id") in collected_doc_ids

    return response_base.success(data=page_data)


@router.post(
    '',
    summary='创建文件',
    dependencies=[
        DependsJwtAuth,
    ],
)
async def create_sys_doc(request: Request, obj: CreateSysDocParam) -> ResponseModel:
    obj.created_by = request.user.id
    obj.created_user = request.user.username
    doc = await sys_doc_service.create(obj=obj)
    await sys_doc_service.create_doc_tokens(id=doc.id)
    return response_base.success(data={'id': doc.id})


@router.put(
    '/{pk}',
    summary='更新文件',
    dependencies=[
        DependsJwtAuth,
    ],
)
async def update_sys_doc(pk: Annotated[int, Path(...)], obj: UpdateSysDocParam) -> ResponseModel:
    count = await sys_doc_service.update(pk=pk, obj=obj)
    await sys_doc_service.create_doc_tokens(id=pk)
    if count > 0:
        return response_base.success()
    return response_base.fail()


@router.delete(
    '',
    summary='（批量）删除文件',
    dependencies=[
        DependsJwtAuth,
    ],
)
async def delete_sys_doc(pk: Annotated[list[int], Query(...)]) -> ResponseModel:
    count = await sys_doc_service.delete(pk=pk)
    if count > 0:
        await sys_doc_service.delete_doc_data(doc_id=pk)
        await sys_doc_service.delete_doc_chunks(doc_id=pk)
        await sys_doc_service.delete_doc_embeddings(doc_id=pk)
        return response_base.success()

    for id in pk:
        doc = await sys_doc_service.get(pk=id)
        if not doc:
            continue
        file = doc.file
        try:
            minio_client.remove_object(bucket_name, file)
        except S3Error:
            continue
    return response_base.fail()
